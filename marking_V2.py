import streamlit as st
import pandas as pd
import requests
import re
import tiktoken
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from io import BytesIO, StringIO
from PyPDF2 import PdfReader

# Configuration
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
ALLOWED_EXTENSIONS = ["docx", "pdf"]
MAX_TOKENS = 7000
PROMPT_BUFFER = 1000

# Initialize encoding
try:
    encoding = tiktoken.encoding_for_model("deepseek-chat")
except KeyError:
    encoding = tiktoken.get_encoding("cl100k_base")

def count_tokens(text):
    return len(encoding.encode(text))

def truncate_text(text, max_tokens):
    tokens = encoding.encode(text)
    return encoding.decode(tokens[:max_tokens]) if len(tokens) > max_tokens else text

def call_deepseek_api(prompt, system_prompt):
    headers = {
        "Authorization": f"Bearer {st.secrets['DEEPSEEK_API_KEY']}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": "deepseek-reasoner",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3,
        "max_tokens": 3000
    }
    
    try:
        response = requests.post(DEEPSEEK_API_URL, json=data, headers=headers)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        st.error(f"API Error: {str(e)}")
        return None

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"DOCX Error: {str(e)}")
        return None

def extract_text_from_pdf(file):
    try:
        reader = PdfReader(file)
        text = []
        for page in reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)
    except Exception as e:
        st.error(f"PDF Error: {str(e)}")
        return None

def parse_csv_section(csv_text):
    try:
        df = pd.read_csv(StringIO(csv_text), quotechar='"', skipinitialspace=True)
        df.columns = df.columns.str.strip().str.lower()
        required_columns = {'criterion', 'score', 'comment'}
        missing = required_columns - set(df.columns)
        if missing:
            st.error(f"Missing columns in CSV: {', '.join(missing)}")
            return None
        df['score'] = pd.to_numeric(df['score'], errors='coerce')
        return df
    except Exception as e:
        st.error(f"CSV Parse Error: {str(e)}")
        return None

def parse_api_response(response):
    try:
        normalized = response.replace('\r\n', '\n')
        
        # Extract sections with improved regex
        csv_match = re.search(
            r'(?i)---CSV_START---(.*?)---CSV_END---', 
            normalized, 
            re.DOTALL
        )
        
        comments_match = re.search(
            r'(?i)---COMMENTS_START---(.*?)---COMMENTS_END---', 
            normalized, 
            re.DOTALL
        )
        
        feedforward_match = re.search(
            r'(?i)---FEEDFORWARD_START---(.*?)---FEEDFORWARD_END---', 
            normalized, 
            re.DOTALL
        )

        if not all([csv_match, comments_match, feedforward_match]):
            raise ValueError("Missing required sections in response")

        def capitalize_sentences(text):
            sentences = re.split(r'(?<=[.!?]) +', text)
            return ' '.join([s[0].upper() + s[1:] for s in sentences if s])

        return {
            'csv': f"Criterion,Score,Comment\n{csv_match.group(1).strip()}",
            'comments': capitalize_sentences(comments_match.group(1).replace('Overall Comments:', '').strip()),
            'feedforward': capitalize_sentences(feedforward_match.group(1).replace('Feedforward:', '').strip())
        }
    except Exception as e:
        st.error(f"Response parsing failed: {str(e)}")
        st.text_area("Raw API Response", response, height=300)
        return None

def extract_weight(criterion_name):
    match = re.search(r'\((\d+)%\)', criterion_name)
    return float(match.group(1)) if match else 0.0

def add_shading(cell):
    shading = OxmlElement('w:shd')
    shading.set(nsdecls('w'), 'fill', 'D9EAD3')
    cell._tc.get_or_add_tcPr().append(shading)

def generate_feedback_doc(student_name, rubric_df, overall_comments, feedforward, total_mark):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    
    # Header with green color
    header = doc.add_heading(f"Feedback for {student_name}", 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    # Rubric Table
    table = doc.add_table(rows=1, cols=len(rubric_df.columns))
    table.style = 'Table Grid'
    
    # Header row with green background
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(rubric_df.columns):
        hdr_cells[i].text = str(col).title()
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        add_shading(hdr_cells[i])
    
    # Data rows with conditional formatting
    for _, row in rubric_df.iterrows():
        row_cells = table.add_row().cells
        for i, (col_name, value) in enumerate(zip(rubric_df.columns, row)):
            cell = row_cells[i]
            cell.text = str(value)
            
            if 'score' in col_name.lower():
                try:
                    score = float(value)
                    for range_col in [col for col in rubric_df.columns if re.match(r'\d+-\d+%', col)]:
                        lower, upper = map(float, range_col.replace('%','').split('-'))
                        if lower <= score <= upper:
                            add_shading(cell)
                            break
                except ValueError:
                    pass
            
            if 'comment' in col_name.lower():
                cell.paragraphs[0].runs[0].text = cell.text.capitalize()
    
    # Feedback sections
    def add_section(heading, content):
        doc.add_heading(heading, level=1).runs[0].font.color.rgb = RGBColor(0, 128, 0)
        doc.add_paragraph(content)
    
    add_section('Overall Comments', overall_comments)
    
    doc.add_heading('Feedforward', level=1).runs[0].font.color.rgb = RGBColor(0, 128, 0)
    for point in feedforward.split('\n'):
        if point.strip():
            p = doc.add_paragraph(style='ListBullet')
            p.add_run(point.strip().lstrip('- '))
    
    total_heading = doc.add_heading('Total Mark', level=1)
    total_heading.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    total_para = doc.add_paragraph()
    total_run = total_para.add_run(f"{total_mark:.2f}%")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(page_title="AutoGrader Pro", layout="wide")
    st.title("✏️ Automated Assignment Grading System")
    
    # Authentication
    if 'authenticated' not in st.session_state:
        password = st.text_input("Enter password:", type="password")
        if st.button("Authenticate") and password == st.secrets["APP_PASSWORD"]:
            st.session_state.authenticated = True
            st.rerun()
        return
    
    st.header("Assignment Configuration")
    assignment_task = st.text_area("Assignment Task & Academic Level", height=150)
    
    st.header("Upload Files")
    # CORRECTED FILE UPLOADER
    rubric_file = st.file_uploader("Rubric (CSV)", type=['csv'])
    submissions = st.file_uploader("Student Submissions", type=ALLOWED_EXTENSIONS, accept_multiple_files=True)
    
    if rubric_file and submissions and st.button("Start Marking"):
        try:
            # Process rubric
            rubric_df = pd.read_csv(rubric_file)
            rubric_df.columns = rubric_df.columns.str.strip().str.lower()
            rubric_df['criterion'] = rubric_df['criterion'].astype(str)
            rubric_df['weight'] = rubric_df['criterion'].apply(extract_weight)
            rubric_df['criterion'] = rubric_df['criterion'].apply(
                lambda x: re.sub(r'\s*\(\d+%\)', '', x).strip().lower()
            )
            
            percentage_columns = [col for col in rubric_df.columns if re.match(r'\d+-\d+%', col)]
            criteria_string = '\n'.join(rubric_df['criterion'].tolist())
            
            for submission in submissions:
                student_name = os.path.splitext(submission.name)[0]
                
                # Extract text
                if submission.type == "application/pdf":
                    text = extract_text_from_pdf(submission)
                else:
                    text = extract_text_from_docx(submission)
                
                if not text:
                    st.error(f"Failed to extract text from {submission.name}")
                    continue
                
                # Enhanced prompt for reference checking
                system_prompt = f"""You are an experienced UK academic. Provide feedback using:
- British English spelling
- Birmingham Newman University guidelines
- Check for reference list existence and Harvard formatting
- Scores between 0-100
- 120 word limits for comments
- Mandatory sections: CSV,Overall Comments, Feedforward"""

                user_prompt = f"""
Generate feedback in EXACTLY this format:

---CSV_START---
Criterion,Score,Comment
"Linking Theory",75,"Good analysis but needs more depth"
"Application of Theory",65,"Adequate but lacks critical engagement"
...
---CSV_END---

---COMMENTS_START---
Overall Comments:
Your essay demonstrates... (150 words max)
---COMMENTS_END---

---FEEDFORWARD_START---
Feedforward:
- Improve critical analysis
- Strengthen theoretical links
- Enhance referencing
---FEEDFORWARD_END---

Submission Content:
{text[:10000]}

Rubric Criteria:
{criteria_string}

Assignment Task:
{assignment_task}
"""
                # API Call
                response = call_deepseek_api(user_prompt, system_prompt)
                if not response:
                    continue
                
                # Parse response
                parsed = parse_api_response(response)
                if not parsed:
                    st.error("Invalid response structure")
                    continue
                
                # Process scores
                scores_df = parse_csv_section(parsed['csv'])
                if scores_df is None or scores_df.empty:
                    st.error("Invalid scores data - check CSV formatting")
                    st.text_area("Raw CSV Data", parsed['csv'], height=200)
                    continue
                
                # Merge dataframes
                try:
                    merged_df = rubric_df.merge(
                        scores_df[['criterion', 'score', 'comment']],
                        on='criterion',
                        how='left'
                    )
                    merged_df['weighted'] = merged_df['weight'] * merged_df['score'] / 100
                    total_mark = merged_df['weighted'].sum()
                except Exception as e:
                    st.error(f"Merge error: {str(e)}")
                    continue
                
                # Generate document
                doc_buffer = generate_feedback_doc(
                    student_name,
                    merged_df[['criterion'] + percentage_columns + ['score', 'comment']],
                    parsed['comments'],
                    parsed['feedforward'],
                    total_mark
                )
                
                st.download_button(
                    label=f"Download {student_name} Feedback",
                    data=doc_buffer,
                    file_name=f"{student_name}_feedback.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Processing Error: {str(e)}")

if __name__ == "__main__":
    main()
    
